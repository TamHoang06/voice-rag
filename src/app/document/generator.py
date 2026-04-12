import os
import io
import base64
import re
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Tuple

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import requests

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image as RLImage
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


class DocumentGenerator:
    """Tạo tài liệu DOCX/PDF có hình ảnh từ văn bản."""
    
    def __init__(self, output_dir: str = "outputs"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        
    # ═══════════════════════════════════════════════════════════════════════
    # DOCX Generation
    # ═══════════════════════════════════════════════════════════════════════
    
    def create_docx(
        self,
        content: str,
        images: Optional[List[str]] = None,
        title: str = "Tài liệu",
        filename: Optional[str] = None
    ) -> str:
        """
        Tạo file DOCX với văn bản và hình ảnh.
        
        Args:
            content: Nội dung văn bản (hỗ trợ markdown đơn giản)
            images: List đường dẫn/URL hình ảnh
            title: Tiêu đề tài liệu
            filename: Tên file output (tự động tạo nếu None)
            
        Returns:
            Đường dẫn file DOCX đã tạo
        """
        doc = Document()
        
        # Tiêu đề
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph()
        
        # Xử lý nội dung (hỗ trợ markdown cơ bản)
        self._add_content_to_docx(doc, content)
        
        # Thêm hình ảnh
        if images:
            doc.add_page_break()
            doc.add_heading("Hình ảnh minh họa", level=1)
            
            for i, img_path in enumerate(images, 1):
                try:
                    img_data = self._load_image(img_path)
                    if img_data:
                        # Lưu tạm để thêm vào docx
                        temp_path = f"temp_img_{i}.png"
                        img_data.save(temp_path)
                        
                        # Thêm caption
                        doc.add_paragraph(f"Hình {i}:", style='Intense Quote')
                        
                        # Thêm ảnh (resize nếu quá lớn)
                        width, height = img_data.size
                        max_width = 6.0  # inches
                        if width > height:
                            doc.add_picture(temp_path, width=Inches(max_width))
                        else:
                            doc.add_picture(temp_path, height=Inches(max_width))
                        
                        doc.add_paragraph()
                        
                        # Xóa file tạm
                        if os.path.exists(temp_path):
                            os.remove(temp_path)
                except Exception as e:
                    print(f"[DOCX] Không thể thêm ảnh {img_path}: {e}")
        
        # Lưu file
        if not filename:
            filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        output_path = os.path.join(self.output_dir, filename)
        doc.save(output_path)
        
        print(f"[DOCX] Saved → {output_path}")
        return output_path
    
    def _add_content_to_docx(self, doc: Document, content: str):
        """Thêm nội dung với format markdown đơn giản."""
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                doc.add_paragraph()
                continue
            
            # Heading
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            
            # Bullet list
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            
            # Numbered list
            elif re.match(r'^\d+\.\s', line):
                doc.add_paragraph(re.sub(r'^\d+\.\s', '', line), style='List Number')
            
            # Normal paragraph
            else:
                p = doc.add_paragraph(line)
                # Bold text
                if '**' in line:
                    p.clear()
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        if i % 2 == 1:  # Odd indices are bold
                            run.bold = True
    
    # ═══════════════════════════════════════════════════════════════════════
    # PDF Generation
    # ═══════════════════════════════════════════════════════════════════════
    
    def create_pdf(
        self,
        content: str,
        images: Optional[List[str]] = None,
        title: str = "Tài liệu",
        filename: Optional[str] = None
    ) -> str:
        """
        Tạo file PDF với văn bản và hình ảnh.
        
        Args:
            content: Nội dung văn bản
            images: List đường dẫn/URL hình ảnh
            title: Tiêu đề tài liệu
            filename: Tên file output
            
        Returns:
            Đường dẫn file PDF đã tạo
        """
        if not filename:
            filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        output_path = os.path.join(self.output_dir, filename)
        
        # Tạo PDF
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        # Tiêu đề
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=RGBColor(0, 51, 102),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Metadata
        meta_style = styles['Normal']
        story.append(Paragraph(
            f"Ngày tạo: {datetime.now().strftime('%d/%m/%Y %H:%M')}",
            meta_style
        ))
        story.append(Spacer(1, 0.3*inch))
        
        # Nội dung
        body_style = ParagraphStyle(
            'BodyText',
            parent=styles['BodyText'],
            fontSize=11,
            alignment=TA_JUSTIFY,
            spaceAfter=12
        )
        
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                # Escape HTML entities
                para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(para, body_style))
                story.append(Spacer(1, 0.1*inch))
        
        # Hình ảnh
        if images:
            story.append(PageBreak())
            story.append(Paragraph("Hình ảnh minh họa", styles['Heading1']))
            story.append(Spacer(1, 0.2*inch))
            
            for i, img_path in enumerate(images, 1):
                try:
                    img_data = self._load_image(img_path)
                    if img_data:
                        # Lưu tạm
                        temp_path = f"temp_pdf_img_{i}.png"
                        img_data.save(temp_path)
                        
                        # Thêm caption
                        story.append(Paragraph(f"<b>Hình {i}:</b>", styles['Normal']))
                        story.append(Spacer(1, 0.1*inch))
                        
                        # Resize ảnh cho vừa trang
                        img = RLImage(temp_path)
                        img.drawHeight = 4*inch
                        img.drawWidth = 4*inch * img.imageWidth / img.imageHeight
                        
                        if img.drawWidth > 6*inch:
                            img.drawWidth = 6*inch
                            img.drawHeight = 6*inch * img.imageHeight / img.imageWidth
                        
                        story.append(img)
                        story.append(Spacer(1, 0.2*inch))
                        
                        # Xóa file tạm
                        if os.path.exists(temp_path):
                            os.remove(temp_path)
                except Exception as e:
                    print(f"[PDF] Không thể thêm ảnh {img_path}: {e}")
        
        # Build PDF
        doc.build(story)
        
        print(f"[PDF] Saved → {output_path}")
        return output_path
    
    # ═══════════════════════════════════════════════════════════════════════
    # Helper Functions
    # ═══════════════════════════════════════════════════════════════════════
    
    def _load_image(self, img_path: str) -> Optional[Image.Image]:
        """Load ảnh từ đường dẫn hoặc URL."""
        try:
            # URL
            if img_path.startswith(('http://', 'https://')):
                response = requests.get(img_path, timeout=10)
                response.raise_for_status()
                return Image.open(io.BytesIO(response.content))
            
            # Base64
            elif img_path.startswith('data:image'):
                base64_data = img_path.split(',')[1]
                return Image.open(io.BytesIO(base64.b64decode(base64_data)))
            
            # File path
            elif os.path.exists(img_path):
                return Image.open(img_path)
            
            else:
                print(f"[IMG] Không tìm thấy: {img_path}")
                return None
                
        except Exception as e:
            print(f"[IMG] Lỗi load ảnh: {e}")
            return None
    
    def extract_images_from_text(self, text: str) -> Tuple[str, List[str]]:
        """
        Tách ảnh từ markdown text.
        
        Returns:
            (clean_text, image_urls)
        """
        # Pattern: ![alt](url) hoặc <img src="url">
        img_pattern = r'!\[.*?\]\((.*?)\)|<img.*?src=["\']?(.*?)["\']?.*?>'
        
        images = []
        for match in re.finditer(img_pattern, text):
            img_url = match.group(1) or match.group(2)
            if img_url:
                images.append(img_url)
        
        # Xóa markdown images khỏi text
        clean_text = re.sub(img_pattern, '', text)
        
        return clean_text, images


# ═══════════════════════════════════════════════════════════════════════════
# Singleton instance
# ═══════════════════════════════════════════════════════════════════════════

_doc_generator = None

def get_document_generator() -> DocumentGenerator:
    """Get singleton instance."""
    global _doc_generator
    if _doc_generator is None:
        _doc_generator = DocumentGenerator()
    return _doc_generator